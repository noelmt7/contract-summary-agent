import io
import pymupdf  # PyMuPDF
import docx
import re
import docx


def extract_tables_from_docx(doc):
    """Extract text from PDF file"""
    tables_text= []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                # Get all text from the cell (including nested tables and paragraphs)
                cell_text = ''
                for paragraph in cell.paragraphs:
                    cell_text += paragraph.text + ' '
                row_data.append(cell_text.strip())
            table_data.append(' | '.join(row_data))
        tables_text.append('\n'.join(table_data))
    return '\n\n'.join(tables_text)

def extract_tables_from_pdf(pdf_document):
    """Extract text and tables from a PDF file."""
    text_content = []

    for page_num in range(len(pdf_document)):
        page = pdf_document[page_num]
        text = page.get_text()

        # Extract tables using PyMuPDF's `find_tables()`
        table_finder = page.find_tables()  # Returns a TableFinder object
        tables = table_finder.tables  # Extract actual table data

        if tables:
            tables_text = []
            for table in tables:
                table_rows = []
                for row in table.extract():  # Extract table contents
                    row_text = ' | '.join(str(cell) for cell in row if cell.strip())
                    if row_text:
                        table_rows.append(row_text)
                if table_rows:
                    tables_text.append('\n'.join(table_rows))

            if tables_text:
                text += '\n\nTABLES:\n' + '\n\n'.join(tables_text)

        text_content.append(text)

    return '\n\n'.join(text_content)


def format_extracted_text(text):
    """Format extracted text by handling bullet points, numbering, and spacing."""
    formatted_text = []
    lines = text.split("\n")

    for line in lines:
        line = line.strip()

        # Convert roman numerals or numbered lists into proper formatting
        if re.match(r"^\(\w+\)", line):  # Matches (i), (ii), (iii), etc.
            line = "\n" + line  # Add spacing before numbered points

        formatted_text.append(line)

    return "\n".join(formatted_text)


def extract_text(file_data, file_type):
    """Extract text from different file formats"""
    try:
        if file_type == 'txt':
            return file_data.decode('utf-8')
        
        
        elif file_type == 'pdf':
            # Use PyMuPDF to extract text and tables
            pdf_document = pymupdf.open(stream=io.BytesIO(file_data), filetype="pdf")
            extracted_text = extract_tables_from_pdf(pdf_document)
            pdf_document.close()
            return extracted_text
        
        elif file_type == 'docx':
            doc = docx.Document(io.BytesIO(file_data))
            
            # Extract regular paragraphs
            paragraphs_text = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract tables
            tables_text = extract_tables_from_docx(doc)
            
            # Combine both with clear separation
            full_text = '\n\n'.join(paragraphs_text)
            if tables_text:
                full_text += '\n\nTABLES:\n' + tables_text
                
            return format_extracted_text(full_text)
        else:
            return "Unsupported file format"
    
    except Exception as e:
        return f"Error extracting text: {str(e)}"

def text_extractor(file_object, file_type):
    """Extract text from uploaded file object"""
    file_bytes = file_object.getvalue()
    extracted_text = extract_text(file_bytes, file_type)
    return extracted_text
